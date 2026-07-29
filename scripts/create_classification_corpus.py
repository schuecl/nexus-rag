"""Create a multi-format corpus spanning every classification and releasability
combination, for exercising the full ingest -> curate -> retrieve path.

`create_upload_test_files.py` answers "does each parser accept its format".
This answers a different question: does the *access-control* path behave
correctly once a realistic amount of text exists at several classification
levels at once. That needs documents that

  - are long enough to chunk more than once (a single-chunk document cannot
    show a ranking or a partial-visibility bug),
  - carry a unique, unguessable canary phrase, so a retrieval result can be
    attributed to exactly one source document rather than matched by filename
    -- which is the bug #226 is about, and
  - cover the classification x releasability matrix rather than sitting at
    UNCLASSIFIED/NONE like the upload-test corpus does.

Everything here is fictional. The domain is deliberately mundane (facilities,
training schedules, equipment logistics) so nothing in the corpus resembles
real controlled information; the classification *markings* are what is under
test, not the content.

Run from the repository root:

    python scripts/create_classification_corpus.py
"""

from __future__ import annotations

import argparse
import datetime as dt
import io
import json
import zipfile
from pathlib import Path

from docx import Document

REPOSITORY_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = REPOSITORY_ROOT / "sample-data" / "classification-corpus"

# Any fixed instant; the value carries no meaning beyond being constant, which
# is what makes regenerating the corpus a no-op in git when nothing changed.
FIXED_TIMESTAMP = dt.datetime(2026, 1, 1, 0, 0, 0, tzinfo=dt.UTC)
FIXED_ZIP_TIME = (2026, 1, 1, 0, 0, 0)

# The ladder and the releasability vocabulary both come from the seeded dev
# realm and the classification_levels table -- see infra/keycloak/realm-export.
# Kept as literals rather than read from the database so the generator runs
# without a stack up; ingest_classification_corpus.py validates them against
# the live vocabulary before uploading anything.
CLASSIFICATIONS = ("UNCLASSIFIED", "CUI", "SECRET")
RELEASABILITIES = ("NONE", "USA", "NOFORN", "NATO", "FVEY")


class Spec:
    """One document in the matrix."""

    def __init__(
        self,
        stem: str,
        fmt: str,
        classification: str,
        releasability: list[str],
        access_scope: list[str],
        title: str,
        canary: str,
        topic: str,
    ) -> None:
        self.stem = stem
        self.fmt = fmt
        self.classification = classification
        self.releasability = releasability
        self.access_scope = access_scope
        self.title = title
        self.canary = canary
        self.topic = topic

    @property
    def filename(self) -> str:
        return f"{self.stem}.{self.fmt}"

    def as_manifest_entry(self) -> dict:
        return {
            "filename": self.filename,
            "classification": self.classification,
            "releasability": self.releasability,
            "access_scope": self.access_scope,
            "canary": self.canary,
            "title": self.title,
        }


# The matrix. Every format appears at every classification, and the
# releasability rotation means no two documents at the same level share a
# holding -- so "which documents can this persona see" has a different answer
# for each seeded user rather than splitting cleanly by classification alone.
SPECS: tuple[Spec, ...] = (
    # -- UNCLASSIFIED -------------------------------------------------------
    Spec(
        "facility-access-hours",
        "txt",
        "UNCLASSIFIED",
        ["NONE"],
        ["ALL_AUTHENTICATED"],
        "Facility Access Hours and Visitor Procedures",
        "MARBLE-HORIZON-7",
        "facility access hours, visitor escort, badge office",
    ),
    Spec(
        "training-calendar-overview",
        "md",
        "UNCLASSIFIED",
        ["USA"],
        ["ALL_AUTHENTICATED"],
        "Annual Training Calendar Overview",
        "COPPER-MERIDIAN-3",
        "training calendar, course scheduling, instructor rotation",
    ),
    Spec(
        "equipment-return-policy",
        "html",
        "UNCLASSIFIED",
        ["NATO"],
        ["ALL_AUTHENTICATED"],
        "Equipment Return and Reconciliation Policy",
        "SLATE-JUNIPER-9",
        "equipment return, inventory reconciliation, custody transfer",
    ),
    Spec(
        "vehicle-dispatch-handbook",
        "pdf",
        "UNCLASSIFIED",
        ["FVEY"],
        ["ALL_AUTHENTICATED"],
        "Vehicle Dispatch Handbook",
        "AMBER-THICKET-2",
        "vehicle dispatch, motor pool, maintenance intervals",
    ),
    Spec(
        "workspace-onboarding-guide",
        "docx",
        "UNCLASSIFIED",
        ["NOFORN"],
        ["ALL_AUTHENTICATED"],
        "Workspace Onboarding Guide",
        "CEDAR-LANTERN-5",
        "onboarding, workstation provisioning, account request",
    ),
    # -- CUI ----------------------------------------------------------------
    Spec(
        "supply-chain-continuity-plan",
        "txt",
        "CUI",
        ["USA"],
        ["ALL_AUTHENTICATED"],
        "Supply Chain Continuity Plan",
        "IRON-PARAPET-4",
        "supply continuity, vendor substitution, stockage objective",
    ),
    Spec(
        "personnel-readiness-summary",
        "md",
        "CUI",
        ["NOFORN"],
        ["ALL_AUTHENTICATED"],
        "Personnel Readiness Summary",
        "QUARTZ-BULWARK-8",
        "readiness reporting, personnel availability, deferral",
    ),
    Spec(
        "network-maintenance-windows",
        "html",
        "CUI",
        ["NATO"],
        ["ALL_AUTHENTICATED"],
        "Network Maintenance Windows",
        "BASALT-KESTREL-1",
        "maintenance window, change freeze, rollback criteria",
    ),
    Spec(
        "facility-power-contingency",
        "pdf",
        "CUI",
        ["FVEY"],
        ["ALL_AUTHENTICATED"],
        "Facility Power Contingency Procedures",
        "ONYX-FURROW-6",
        "power contingency, generator load, transfer switch",
    ),
    Spec(
        "records-retention-schedule",
        "docx",
        "CUI",
        ["NONE"],
        ["ALL_AUTHENTICATED"],
        "Records Retention Schedule",
        "PEWTER-GABLE-0",
        "retention schedule, disposition authority, destruction record",
    ),
    # -- SECRET -------------------------------------------------------------
    Spec(
        "exercise-participation-roster",
        "txt",
        "SECRET",
        ["NOFORN"],
        ["ALL_AUTHENTICATED"],
        "Exercise Participation Roster Procedures",
        "GRANITE-VESPER-2",
        "exercise roster, participation control, roster amendment",
    ),
    Spec(
        "site-survey-findings",
        "md",
        "SECRET",
        ["FVEY"],
        ["ALL_AUTHENTICATED"],
        "Site Survey Findings and Recommendations",
        "COBALT-TRELLIS-7",
        "site survey, siting constraint, remediation sequence",
    ),
    Spec(
        "contingency-communications-plan",
        "html",
        "SECRET",
        ["NATO"],
        ["ALL_AUTHENTICATED"],
        "Contingency Communications Plan",
        "TUNGSTEN-ORCHARD-4",
        "contingency communications, relay plan, callsign rotation",
    ),
    Spec(
        "logistics-priority-matrix",
        "pdf",
        "SECRET",
        ["USA"],
        ["ALL_AUTHENTICATED"],
        "Logistics Priority Matrix",
        "SILVER-RAMPART-9",
        "logistics priority, allocation tier, expedite authority",
    ),
    Spec(
        "installation-security-review",
        "docx",
        "SECRET",
        ["NONE"],
        ["ALL_AUTHENTICATED"],
        "Installation Security Review",
        "NICKEL-PALISADE-3",
        "security review, access control finding, corrective action",
    ),
)


def _paragraphs(spec: Spec) -> list[tuple[str, list[str]]]:
    """Body content for one document, as (heading, paragraphs) sections.

    Long enough to chunk several times, and every section repeats the canary
    so a retrieval hit anywhere in the document is attributable. The facts are
    invented and internally consistent per document, which is what makes a
    golden query answerable rather than merely retrievable.
    """
    c = spec.canary
    subject = spec.topic.split(",")[0].strip()
    return [
        (
            "Purpose and Scope",
            [
                f"This document establishes the procedures governing {spec.topic} for "
                f"the fictional Northwind Support Activity. It is issued under reference "
                f"code {c} and supersedes all prior guidance carrying that reference. "
                f"Every procedure below applies to assigned, attached, and visiting "
                f"personnel unless a paragraph states otherwise.",
                f"The scope of {c} covers routine operations only. Contingency operations "
                f"are governed separately and are outside the scope of this document. "
                f"Where this document and a contingency instruction disagree, the "
                f"contingency instruction governs for the duration of the contingency "
                f"and this document resumes effect on termination.",
                f"Questions about the application of {c} go to the issuing office. The "
                f"issuing office maintains the authoritative copy; printed copies are "
                f"reference only and are not controlled once printed.",
            ],
        ),
        (
            "Responsibilities",
            [
                f"The issuing office owns {c} and reviews it annually, or sooner when a "
                f"referenced authority changes. Review completion is recorded in the "
                f"document control log with the reviewer's name and the review date.",
                f"Section supervisors implement {c} within their sections, confirm that "
                f"assigned personnel have read it, and forward proposed changes through "
                f"the document control process rather than applying local variations. A "
                f"local variation that is not recorded is a finding.",
                f"Individual personnel comply with {c}, report conditions that prevent "
                f"compliance, and do not defer a required step on the assumption that "
                f"someone else has completed it. Deferral without a recorded handoff is "
                f"the single most common cause of the discrepancies discussed below.",
            ],
        ),
        (
            "Procedures",
            [
                f"Step one: verify the current state before making any change. For "
                f"{subject}, this means confirming the record matches the physical "
                f"condition and noting any discrepancy before proceeding. Under {c} a "
                f"discrepancy noted before the change is a correction; the same "
                f"discrepancy noted afterward is an incident.",
                f"Step two: perform the action in the sequence given. The sequence in {c} "
                f"is not arbitrary -- several steps establish a condition that a later "
                f"step depends on, and performing them out of order produces a result "
                f"that appears correct while leaving the underlying condition unmet.",
                f"Step three: record the outcome the same working day. A record entered "
                f"later than the same working day is annotated as a late entry. Late "
                f"entries are permitted but are counted, and a section exceeding four "
                f"late entries in a quarter is reviewed under {c}.",
                f"Step four: confirm the record. Confirmation is a separate action by a "
                f"second person, and it is the step most often skipped when a section is "
                f"short-handed. {c} does not authorize self-confirmation under any "
                f"staffing condition.",
            ],
        ),
        (
            "Common Discrepancies",
            [
                f"The most frequent discrepancy found against {c} is a record that was "
                f"opened and never closed. The action was completed, the outcome was "
                f"correct, and nobody returned to the record. This is treated as an "
                f"administrative finding rather than an operational one, but repeated "
                f"instances in one section are treated as a process failure.",
                f"The second most frequent is a transfer of responsibility that was "
                f"verbal. Under {c} a transfer is effective when it is recorded, not when "
                f"it is spoken, and the person named in the record remains responsible "
                f"until a replacement record exists.",
                f"The third is the use of a superseded copy. Because printed copies are "
                f"not controlled, a printed {c} more than thirty days old should be "
                f"discarded and reprinted rather than annotated by hand.",
            ],
        ),
        (
            "Review and Records",
            [
                f"Records generated under {c} are retained for three years from the date "
                f"of the entry, then dispositioned in accordance with the applicable "
                f"retention schedule. Retention runs from the entry date, not from the "
                f"date the record was closed.",
                f"The annual review of {c} produces one of three outcomes: reissued "
                f"without change, reissued with change, or cancelled. All three are "
                f"recorded; a review that produces no record is not a review.",
                f"This document, {c}, takes effect on issue and remains in effect until "
                f"cancelled or superseded by a document that names it explicitly.",
            ],
        ),
    ]


def _plain_lines(spec: Spec) -> list[str]:
    lines = [spec.title, "=" * len(spec.title), ""]
    for heading, paras in _paragraphs(spec):
        lines.extend([heading, "-" * len(heading), ""])
        for para in paras:
            lines.extend([para, ""])
    return lines


def write_txt(spec: Spec, path: Path) -> None:
    path.write_text("\n".join(_plain_lines(spec)), encoding="utf-8")


def write_md(spec: Spec, path: Path) -> None:
    out = [f"# {spec.title}", ""]
    for heading, paras in _paragraphs(spec):
        out.extend([f"## {heading}", ""])
        for para in paras:
            out.extend([para, ""])
    path.write_text("\n".join(out), encoding="utf-8")


def _html_escape(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def write_html(spec: Spec, path: Path) -> None:
    out = [
        "<!doctype html>",
        '<html lang="en">',
        "<head>",
        '<meta charset="utf-8" />',
        f"<title>{_html_escape(spec.title)}</title>",
        "</head>",
        "<body>",
        f"<h1>{_html_escape(spec.title)}</h1>",
    ]
    for heading, paras in _paragraphs(spec):
        out.append(f"<h2>{_html_escape(heading)}</h2>")
        out.extend(f"<p>{_html_escape(para)}</p>" for para in paras)
    out.extend(["</body>", "</html>"])
    path.write_text("\n".join(out), encoding="utf-8")


def write_docx(spec: Spec, path: Path) -> None:
    document = Document()
    document.add_heading(spec.title, level=1)
    for heading, paras in _paragraphs(spec):
        document.add_heading(heading, level=2)
        for para in paras:
            document.add_paragraph(para)

    # python-docx stamps "now" into the core properties, and zipfile stamps it
    # again into every entry header, so two runs of this generator produce two
    # different files with identical content. That defeats the point of a
    # committed corpus: `git status` reports a diff after any regeneration, and
    # a reviewer cannot tell a real content change from a re-run. Both clocks
    # are pinned to a fixed instant instead.
    document.core_properties.created = FIXED_TIMESTAMP
    document.core_properties.modified = FIXED_TIMESTAMP
    document.core_properties.last_modified_by = "nexus-rag corpus generator"
    document.core_properties.revision = 1

    buffer = io.BytesIO()
    document.save(buffer)
    _rewrite_zip_deterministically(buffer.getvalue(), path)


def _rewrite_zip_deterministically(data: bytes, path: Path) -> None:
    """Repack an OOXML file with fixed entry timestamps and stable ordering."""
    with zipfile.ZipFile(io.BytesIO(data)) as source:
        entries = sorted(source.infolist(), key=lambda i: i.filename)
        payloads = [(info.filename, source.read(info.filename)) for info in entries]
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as target:
        for name, payload in payloads:
            info = zipfile.ZipInfo(name, date_time=FIXED_ZIP_TIME)
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = 0o600 << 16
            target.writestr(info, payload)


def _pdf_escape(text: str) -> str:
    return text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def _wrap(text: str, width: int = 92) -> list[str]:
    words, line, out = text.split(), "", []
    for word in words:
        if line and len(line) + 1 + len(word) > width:
            out.append(line)
            line = word
        else:
            line = f"{line} {word}".strip()
    if line:
        out.append(line)
    return out


def write_pdf(spec: Spec, path: Path) -> None:
    """Write a multi-page PDF without adding a PDF dependency.

    Same approach as create_upload_test_files.py's single-page writer, extended
    to paginate -- these documents are far too long for one page, and a
    one-page PDF would not exercise the parser's page handling at all.
    """
    lines: list[tuple[str, bool]] = [(spec.title, True)]
    for heading, paras in _paragraphs(spec):
        lines.append(("", False))
        lines.append((heading, True))
        for para in paras:
            lines.extend((wrapped, False) for wrapped in _wrap(para))
            lines.append(("", False))

    lines_per_page = 46
    pages = [lines[i : i + lines_per_page] for i in range(0, len(lines), lines_per_page)]

    streams: list[bytes] = []
    for page_lines in pages:
        commands = ["BT", "/F1 11 Tf", "14 TL", "54 738 Td"]
        for text, bold in page_lines:
            commands.append("/F2 12 Tf" if bold else "/F1 10 Tf")
            commands.append(f"({_pdf_escape(text)}) Tj" if text else "()  Tj")
            commands.append("T*")
        commands.append("ET")
        streams.append("\n".join(commands).encode("ascii", "replace"))

    # Object layout: 1 catalog, 2 pages, 3 font F1, 4 font F2, then per page a
    # page object followed by its content stream.
    n_pages = len(pages)
    first_page_obj = 5
    kids = " ".join(f"{first_page_obj + 2 * i} 0 R" for i in range(n_pages))

    objects: list[bytes] = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        f"<< /Type /Pages /Kids [{kids}] /Count {n_pages} >>".encode("ascii"),
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica-Bold >>",
    ]
    for index, stream in enumerate(streams):
        content_obj = first_page_obj + 2 * index + 1
        objects.append(
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 3 0 R /F2 4 0 R >> >> "
            + f"/Contents {content_obj} 0 R >>".encode("ascii")
        )
        objects.append(
            f"<< /Length {len(stream)} >>\nstream\n".encode("ascii") + stream + b"\nendstream"
        )

    out = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for number, body in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{number} 0 obj\n".encode("ascii") + body + b"\nendobj\n"
    xref_at = len(out)
    out += f"xref\n0 {len(objects) + 1}\n".encode("ascii")
    out += b"0000000000 65535 f \n"
    for offset in offsets[1:]:
        out += f"{offset:010d} 00000 n \n".encode("ascii")
    out += (
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_at}\n%%EOF\n"
    ).encode("ascii")
    path.write_bytes(bytes(out))


WRITERS = {
    "txt": write_txt,
    "md": write_md,
    "html": write_html,
    "pdf": write_pdf,
    "docx": write_docx,
}


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--output-dir", type=Path, default=OUTPUT_DIR)
    args = parser.parse_args()

    args.output_dir.mkdir(parents=True, exist_ok=True)
    total = 0
    for spec in SPECS:
        path = args.output_dir / spec.filename
        WRITERS[spec.fmt](spec, path)
        total += path.stat().st_size
        print(
            f"  {spec.filename:38s} {spec.classification:13s} "
            f"{'/'.join(spec.releasability):6s} {path.stat().st_size:7d} bytes"
        )

    manifest = args.output_dir / "manifest.json"
    manifest.write_text(
        json.dumps([spec.as_manifest_entry() for spec in SPECS], indent=2) + "\n",
        encoding="utf-8",
    )
    print(f"\n{len(SPECS)} documents, {total} bytes total")
    print(f"manifest: {manifest.relative_to(REPOSITORY_ROOT)}")


if __name__ == "__main__":
    main()
